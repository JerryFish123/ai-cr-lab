"""Download PRD attachments and extract plain text from PDF / Word.

Downloaded binaries are written under a temp dir and **always deleted** after
text extraction (success or failure), so PRD originals do not accumulate on disk.
"""
from __future__ import annotations

import io
import os
import tempfile
import time
from dataclasses import dataclass
from pathlib import Path
from urllib.parse import urlparse

import requests

from biz.utils.log import logger

# Soft cap to keep agent prompts bounded (chars, not tokens).
DEFAULT_MAX_CHARS = int(os.getenv("PRD_TEXT_MAX_CHARS", "50000"))
# Max bytes to accept for a single PRD download (default 20 MiB).
DEFAULT_MAX_DOWNLOAD_BYTES = int(os.getenv("PRD_MAX_DOWNLOAD_BYTES", str(20 * 1024 * 1024)))


@dataclass
class ExtractResult:
    ok: bool
    text: str = ""
    reason: str = ""
    source_url: str = ""
    content_type: str = ""


def prd_tmp_root() -> Path:
    """Directory for short-lived PRD downloads (gitignored)."""
    root = Path(os.getenv("PRD_TMP_DIR", "data/prd_tmp"))
    root.mkdir(parents=True, exist_ok=True)
    return root


def safe_unlink(path: Path | str | None) -> None:
    """Best-effort delete; never raise to callers."""
    if not path:
        return
    p = Path(path)
    try:
        p.unlink(missing_ok=True)
    except TypeError:
        # Python <3.8 compatibility not needed; still guard odd path types
        try:
            if p.exists():
                p.unlink()
        except Exception as e:  # noqa: BLE001
            logger.warning("failed to delete PRD temp file %s: %s", p, e)
    except Exception as e:  # noqa: BLE001
        logger.warning("failed to delete PRD temp file %s: %s", p, e)


def cleanup_stale_prd_temps(*, max_age_seconds: int | None = None) -> int:
    """Remove leftover temp files from crashed runs. Returns deleted count."""
    age = max_age_seconds
    if age is None:
        age = int(os.getenv("PRD_TMP_MAX_AGE_SECONDS", "3600"))
    root = prd_tmp_root()
    now = time.time()
    deleted = 0
    for p in root.glob("prd_*"):
        try:
            if not p.is_file():
                continue
            if now - p.stat().st_mtime > age:
                p.unlink(missing_ok=True)
                deleted += 1
        except Exception as e:  # noqa: BLE001
            logger.warning("stale PRD temp cleanup failed for %s: %s", p, e)
    return deleted


def _guess_kind(url: str, content_type: str, magic: bytes) -> str:
    lower = (url or "").lower()
    ct = (content_type or "").lower()
    # Images (common for GitHub user-attachments/assets without extension)
    if magic[:3] == b"\xff\xd8\xff" or magic[:8] == b"\x89PNG\r\n\x1a\n" or magic[:4] == b"GIF8":
        return "image"
    if "image/" in ct:
        return "image"
    if lower.endswith(".pdf") or "application/pdf" in ct or magic.startswith(b"%PDF"):
        return "pdf"
    if lower.endswith(".docx") or "wordprocessingml" in ct or magic[:2] == b"PK":
        # docx is a zip; .doc is OLE — distinguish by extension / content-type
        if lower.endswith(".doc") and not lower.endswith(".docx"):
            return "doc"
        if "msword" in ct and "openxml" not in ct:
            return "doc"
        return "docx"
    if lower.endswith(".doc") or magic[:4] == b"\xd0\xcf\x11\xe0":
        return "doc"
    return "unknown"


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001 — keep going across pages
            logger.warning("pdf page extract failed: %s", e)
    return "\n".join(parts).strip()


def _extract_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()


def extract_bytes(data: bytes, *, url: str = "", content_type: str = "") -> ExtractResult:
    """Extract text from in-memory file bytes."""
    if not data:
        return ExtractResult(ok=False, reason="下载内容为空", source_url=url, content_type=content_type)

    kind = _guess_kind(url, content_type, data[:8])
    try:
        if kind == "pdf":
            text = _extract_pdf(data)
        elif kind == "docx":
            text = _extract_docx(data)
        elif kind == "doc":
            return ExtractResult(
                ok=False,
                reason="暂不支持旧版 .doc（Word 97-2003），请上传 .docx 或 PDF",
                source_url=url,
                content_type=content_type,
            )
        elif kind == "image":
            return ExtractResult(
                ok=False,
                reason="下载内容是图片而非 PDF/Word（GitHub assets 链常无扩展名；请确认拖入的是 PRD 文档）",
                source_url=url,
                content_type=content_type,
            )
        else:
            return ExtractResult(
                ok=False,
                reason=f"无法识别文件类型（url/content-type/magic 均未匹配 PDF/DOCX）: content_type={content_type!r}",
                source_url=url,
                content_type=content_type,
            )
    except Exception as e:  # noqa: BLE001
        return ExtractResult(
            ok=False,
            reason=f"文本抽取异常: {type(e).__name__}: {e}",
            source_url=url,
            content_type=content_type,
        )

    if not text or not text.strip():
        return ExtractResult(
            ok=False,
            reason="抽取出的文本为空（可能是扫描版 PDF 或受保护文档）",
            source_url=url,
            content_type=content_type,
        )

    clipped = text.strip()
    if len(clipped) > DEFAULT_MAX_CHARS:
        clipped = clipped[:DEFAULT_MAX_CHARS] + f"\n\n…(已截断，原文约 {len(text)} 字)"

    return ExtractResult(ok=True, text=clipped, source_url=url, content_type=content_type)


def _stream_download_to_temp(
    url: str,
    headers: dict,
    *,
    timeout: int,
    max_bytes: int,
) -> tuple[Path | None, int, str, str | None]:
    """Stream response body to a temp file.

    Returns (tmp_path, status_code, content_type, error_reason).
    On error, tmp_path may still be set (caller must delete).
    """
    tmp_path: Path | None = None
    fd, name = tempfile.mkstemp(prefix="prd_", suffix=".bin", dir=str(prd_tmp_root()))
    os.close(fd)
    tmp_path = Path(name)

    with requests.get(url, headers=headers, timeout=timeout, allow_redirects=True, stream=True) as resp:
        ctype = resp.headers.get("Content-Type", "")
        if resp.status_code != 200:
            return tmp_path, resp.status_code, ctype, (
                f"下载失败: HTTP {resp.status_code}（若为 GitHub 附件，请确认 Token 有权访问）"
            )

        written = 0
        with open(tmp_path, "wb") as f:
            for chunk in resp.iter_content(chunk_size=64 * 1024):
                if not chunk:
                    continue
                written += len(chunk)
                if written > max_bytes:
                    return tmp_path, resp.status_code, ctype, (
                        f"下载超过大小上限 {max_bytes} bytes，已中止"
                    )
                f.write(chunk)

        return tmp_path, resp.status_code, ctype, None


def _is_github_host_url(url: str) -> bool:
    host = (urlparse(url).hostname or "").lower()
    return host == "github.com" or host.endswith(".github.com") or host.endswith(
        "githubusercontent.com"
    )


def _download_timeout(url: str, timeout: int | tuple[int, int]) -> int | tuple[int, int]:
    """Use a shorter read timeout for github.com hosts (often unreachable in CN)."""
    if isinstance(timeout, tuple):
        return timeout
    if _is_github_host_url(url):
        connect = int(os.getenv("PRD_GITHUB_CONNECT_TIMEOUT", "10"))
        read = int(os.getenv("PRD_GITHUB_READ_TIMEOUT", "20"))
        return (connect, read)
    return timeout


def download_and_extract(
    url: str,
    token: str | None = None,
    *,
    timeout: int | tuple[int, int] = 60,
) -> ExtractResult:
    """Download attachment to a temp file, extract text, then delete the file.

    Guarantees best-effort cleanup in ``finally`` so PRD binaries do not linger.
    """
    if not url:
        return ExtractResult(ok=False, reason="附件 URL 为空")

    # Opportunistic cleanup of orphans from previous crashed runs.
    try:
        cleanup_stale_prd_temps()
    except Exception as e:  # noqa: BLE001
        logger.warning("prd stale temp cleanup skipped: %s", e)

    headers = {
        "User-Agent": "ai-cr-lab-prd-extractor/1.0",
        "Accept": "*/*",
    }
    if token:
        headers["Authorization"] = f"Bearer {token}"

    effective_timeout = _download_timeout(url, timeout)
    tmp_path: Path | None = None
    try:
        try:
            tmp_path, status, ctype, err = _stream_download_to_temp(
                url,
                headers,
                timeout=effective_timeout,
                max_bytes=DEFAULT_MAX_DOWNLOAD_BYTES,
            )
            if status in (401, 403) and token:
                safe_unlink(tmp_path)
                tmp_path = None
                headers["Authorization"] = f"token {token}"
                tmp_path, status, ctype, err = _stream_download_to_temp(
                    url,
                    headers,
                    timeout=effective_timeout,
                    max_bytes=DEFAULT_MAX_DOWNLOAD_BYTES,
                )
        except requests.RequestException as e:
            return ExtractResult(ok=False, reason=f"下载失败: {type(e).__name__}: {e}", source_url=url)

        if err:
            return ExtractResult(ok=False, reason=err, source_url=url, content_type=ctype)

        assert tmp_path is not None
        data = tmp_path.read_bytes()

        # Reject HTML login / error pages mistaken as files
        if "text/html" in (ctype or "").lower():
            head = data[:200].lower()
            if not url.lower().endswith((".pdf", ".docx", ".doc")) or b"<html" in head:
                return ExtractResult(
                    ok=False,
                    reason=f"下载结果为 HTML 而非文档（可能未鉴权）: Content-Type={ctype}",
                    source_url=url,
                    content_type=ctype,
                )

        return extract_bytes(data, url=url, content_type=ctype)
    finally:
        safe_unlink(tmp_path)


def resolve_and_extract_prd(
    url: str,
    token: str | None = None,
    *,
    repo_key: str | None = None,
    ref: str | None = None,
    timeout: int | tuple[int, int] = 60,
) -> ExtractResult:
    """Resolve PRD text: GitHub Contents API / HTTP download / repo PRD fallback.

    Order:
      1. Prefer Contents API for blob/raw URLs, and for user-attachments when
         repo_key/ref are known (avoids hanging github.com on domestic ECS)
      2. Direct HTTP download of the attachment URL
      3. On failure (non-API-first path): search repo tree for a PRD pdf/docx
    """
    from biz.prd.github_prd_source import (
        parse_github_repo_file_url,
        resolve_prd_bytes_via_github_api,
    )

    if not url:
        return ExtractResult(ok=False, reason="附件 URL 为空")

    prefer_api_first = bool(
        token
        and (
            parse_github_repo_file_url(url)
            or (
                repo_key
                and ref
                and "user-attachments/" in url.lower()
            )
        )
    )
    if prefer_api_first:
        data, source = resolve_prd_bytes_via_github_api(
            url=url, token=token, repo_key=repo_key, ref=ref
        )
        if data:
            result = extract_bytes(data, url=url)
            if result.ok:
                logger.info("PRD resolved via %s", source)
                return result
            return result

    http_result = download_and_extract(url, token, timeout=timeout)
    if http_result.ok:
        return http_result

    if token and repo_key and ref and not prefer_api_first:
        data, source = resolve_prd_bytes_via_github_api(
            url=url, token=token, repo_key=repo_key, ref=ref
        )
        if data:
            result = extract_bytes(data, url=url)
            if result.ok:
                logger.info(
                    "PRD HTTP failed (%s); recovered via %s",
                    http_result.reason,
                    source,
                )
                return result
            return result
        return ExtractResult(
            ok=False,
            reason=f"{http_result.reason}；{source}",
            source_url=url,
        )

    if prefer_api_first:
        return ExtractResult(
            ok=False,
            reason=f"Contents API 与直连下载均失败: {http_result.reason}",
            source_url=url,
        )

    return http_result


def extract_local_file(path: str | Path) -> ExtractResult:
    """Helper for tests: extract from a local path (does not delete the source)."""
    p = Path(path)
    if not p.is_file():
        return ExtractResult(ok=False, reason=f"文件不存在: {p}")
    data = p.read_bytes()
    return extract_bytes(data, url=str(p), content_type="")
